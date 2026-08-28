from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn


def create_word_document(
    title: str,
    paragraphs: list[str],
    output_dir: Path,
) -> str:
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(11)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(20)

    document.add_paragraph()

    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.add_run(text)

    document.add_paragraph()

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run(
        f"由 AI 办公工具生成 · {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    footer_run.font.size = Pt(9)
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = output_dir / filename

    document.save(output_path)

    return filename